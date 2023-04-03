import re
import os
import selenium
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
import time
import datetime
import pandas as pd
from alive_progress import alive_bar
from PIL import Image
import pyautogui as p
import glob
import docx
from docx.shared import Inches, Pt, Cm

device = int(input("\n Type:\n 1 for Laptop \n 2 for PC\n"))
if device not in range(1,3):
    print('\nInvalid input')
    exit()

options = Options()
options.add_experimental_option('detach', True)
username_array = os.getlogin()
try:
    options.add_argument(r"--user-data-dir=C:\Users\{}\AppData\Local\Google\Chrome\User Data\Default".format(username_array))
except:
    pass
options.add_argument("--start-maximized")
options.add_argument('ignore-certificate-errors')
driver = webdriver.Chrome(chrome_options=options)
action = ActionChains(driver)

links = [
    'https://10.37.155.49/dashboard',
    'https://10.37.155.88/dashboard',
    'https://10.37.155.93/dashboard',
    'https://10.37.155.98/dashboard',
    'https://10.37.155.103/dashboard',
    'https://10.37.155.108/dashboard',
    'https://10.37.155.53/dashboard',
    'https://10.37.155.113/dashboard',
    'https://10.37.155.61/dashboard',  
    'https://10.51.123.50/dashboard',
    'https://10.51.123.56/',
    'https://10.51.123.112/dashboard',
    'https://10.51.123.117/dashboard',
    'https://10.51.123.122/dashboard',
    'https://10.51.123.127/dashboard',
    'https://10.51.123.68/dashboard',
    'https://10.51.123.132/',
    'https://10.51.123.76/dashboard'  
]

arrays = [
    "FAL-PURE-X70-01",
    "FAL-Pure-X90-01",
    "FAL-Pure-X90-02",
    "FAL-Pure-X90-03",
    "FAL-Pure-X90-04",
    "FAL-Pure-X90-05",
    "FAL-Pure-C60-01",
    "FAL-Pure-C60-02",
    "FAL-Pure-FB-01" ,
    "BDB-Pure-FA-X70",
    "BDB-Pure-X70-01",
    "BDB-Pure-X90-01",
    "BDB-Pure-X90-02",
    "BDB-Pure-X90-03",
    "BDB-Pure-X90-04",
    "BDB-Pure-C60-01",
    "BDB-Pure-C60-02",
    "BDB-Pure-FB-01"
]

username_array = **
username_pure1 = **
psw = **

purity_os_ls, usage_value_ls, data_reduction_ls, used_ls, provisioned_ls, system_ls, array_ls = [],[],[],[],[],[],[]

def main():
    loop_through_links()
    data_to_df()
    write_to_csv()
    screenshot()
    driver.quit()

def login_array():
    # Clicking prereqs to login_array if any (advance, proceed, etc)
    try:
        driver.find_element(By.XPATH, '//*[@id="details-button"]').click()
        driver.find_element(By.XPATH, '//*[@id="proceed-link"]').click()
    except:
        pass

    # Logging in
    try:
        driver.find_element(By.XPATH, '//*[@id="username"]').send_keys(username_array)
        driver.find_element(By.XPATH, '//*[@id="password"]').send_keys(psw)
        driver.find_element(By.XPATH, '//*[@id="password"]').send_keys(Keys.RETURN)     
    except:
        pass


def login_pure1():
    # Logging in
    try:
        WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.form-control'))).send_keys(username_pure1)
        driver.find_element(By.CSS_SELECTOR, '.form-control').send_keys(Keys.RETURN)  
        driver.find_element(By.XPATH, '//*[@id="loginPage:loginForm"]/div[1]/div/div/input[2]').send_keys(psw)
        driver.find_element(By.XPATH, '//*[@id="loginPage:loginForm"]/div[1]/div/div/input[2]').send_keys(Keys.RETURN)     
    except:
        pass


def data_collection():
    global system, system_label, usage_value, data_reduction, used, provisioned, purity_os, array,ls
    
    
    data_reduction = WebDriverWait(driver, 40).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, '.capacity-value')))[0].text
    data_reduction = remove_extra_words(data_reduction)
    
    used = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="tab-content"]/div/div/dashboard/div/div[1]/dashboard-capacity-details/div/div[2]/div[2]/div[2]/div[3]/div[1]/div[1]'))).text
    used = remove_extra_words(used)
   
    array = WebDriverWait(driver, 5).until(EC.visibility_of_element_located((By.CSS_SELECTOR, '.sidebar-info [id="sidebar-array-name"]'))).text
    
    try:
        system = WebDriverWait(driver, 5).until(EC.visibility_of_all_elements_located((By.CLASS_NAME, 'legend-value-container')))[0].text
    except:
        pass
    
    try:
        usage_value = WebDriverWait(driver, 5).until(EC.visibility_of_element_located((By.CSS_SELECTOR, '.percent-used .value'))).text
    except:
        usage_value = WebDriverWait(driver, 5).until(EC.visibility_of_element_located((By.CLASS_NAME, 'value'))).text
    
    try:
        provisioned = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="tab-content"]/div/div/dashboard/div/div[1]/dashboard-capacity-details/div/div[2]/div[2]/div[3]/div[4]'))).text
        provisioned = remove_extra_words(provisioned)
    except:
        provisioned = '-'
        
    try:    
        purity_os = WebDriverWait(driver, 5).until(EC.visibility_of_all_elements_located((By.XPATH, '//*[@id="sidebar-array-version"]/strong'))).text
    except:
        purity_os = '-'
    
    
    ls = [data_reduction, used, array, system,usage_value,provisioned,purity_os]
    # remove_extra_words()
    # print(new_ls)
    
    list_converter()

    
def remove_extra_words(data):
    global sub_string

    match = re.search(r'(\d+[.]?[\d+]?)', data)
    if match:
        sub_string = match.group(0)
    return(sub_string)
  
    
def loop_through_links():
    with alive_bar(100) as bar:  
        for link in links:
            driver.get(link)
            login_array()
            data_collection()
            bar()                  
        
        
def list_converter():  
    global purity_os_ls,usage_value_ls,data_reduction_ls,used_ls,provisioned_ls,system_ls,array_ls
      
    purity_os_ls.append(purity_os)
    usage_value_ls.append(usage_value)
    data_reduction_ls.append(data_reduction)
    used_ls.append(used)
    provisioned_ls.append(provisioned)
    system_ls.append(system)
    array_ls.append(array)
    
    
def data_to_df():
    global df
    
    data_dict = {
        'Array': array_ls,
        'Purity version': purity_os_ls,
        'Usage': usage_value_ls,
        'Data Reduction': data_reduction_ls,
        'used': used_ls,
        'Provisioned': provisioned_ls,
        'System': system_ls
    }
    df = pd.DataFrame(data_dict)
    
    
def write_to_csv():
    now = datetime.datetime.now()
    df.to_csv(r'D:\Temp\VM Pure Array Health check\{}.csv'.format(now.strftime("%d-%m-%Y")), encoding='utf-8')


def screenshot():
    global now
    
    driver.get("https://pure1.purestorage.com/analysis/performance/arrays/array?timeRange=%5B%7B%22key%22%3A%22timeIdx%22%2C%22value%22%3A%222%22%7D%5D&selection=eb9d39f7-d1be-47f3-a414-fec55d83a63c")
    login_pure1()
    pure1_prereqs()
    now = datetime.datetime.now()
    
    files = glob.glob(r'D:\Temp\VM Pure Array Health check\Screenshots\*') #removing any file inside images folder
    for f in files:
        os.remove(f)
    
    dimensions_for_pc()

    
    for array in arrays:
        search.clear()
        search.send_keys(array)
        search.send_keys(Keys.RETURN)
        WebDriverWait(driver, 20).until(EC.invisibility_of_element_located((By.CSS_SELECTOR, '.is-loading')))
        time.sleep(1)
        
        checkbox.click() # Ticking

        if device == 1:
            dimensions_for_laptop()
        else:
            dimensions_for_pc()
        
        WebDriverWait(driver, 120).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, '.highcharts-series')))
        WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.performance-chart-container')))
        
        driver.execute_script("window.scrollTo(0,document.body.scrollHeight)")
        
        ss = p.screenshot(region=(x,y,w,h))
        ss.save(r'D:\Temp\VM Pure Array Health check\Screenshots\{} [{}].png'.format(array, now.strftime("%d-%m-%Y")))
        
        driver.find_element(By.CSS_SELECTOR, '.entity-selection-button .selection-button').click()
        clear_selections = driver.find_element(By.CSS_SELECTOR, '.clear-all').click()
        
    saving_ss_into_word()

def dimensions_for_pc():
    global x,y,w,h
    
    x = 397
    y = 463
    w = 1612 - 397
    h = 1006 - 463
    
def dimensions_for_laptop():
    global x,y,w,h
    
    x = 330
    y = 205
    w = 1290 - 360
    h = 680 - 205

def pure1_prereqs():
    # Setting the page up accurately for screenshot
    global search, selections,chart,checkbox,clear_selections
    # chart_loader = WebDriverWait(driver, 20).until(EC.staleness_of((By.CSS_SELECTOR, '.highcharts-loading-inner')))
    WebDriverWait(driver, 120).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, '.highcharts-series')))
    WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.performance-chart-container')))
        
    # Scrolling down
    driver.execute_script("document.body.style.transform='scale(0.7)';")
    chart = driver.find_element(By.CSS_SELECTOR, '.performance-chart-container')
    
    checkbox = driver.find_element(By.CSS_SELECTOR, '.select-all-check')
    search = driver.find_element(By.CSS_SELECTOR, '.form-control')
    selections = driver.find_element(By.CSS_SELECTOR, '.entity-selection-button .selection-button').text
    
    if '(0)' not in str(selections):
        driver.find_element(By.CSS_SELECTOR, '.entity-selection-button .selection-button').click()
        clear_selections = driver.find_element(By.CSS_SELECTOR, '.clear-all').click()
    
    

def saving_ss_into_word():
    ## Making a word file
    # editing margins
    doc = docx.Document() # opening doc
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
        
    # heading 
    para = doc.add_paragraph()
    title = para.add_run('DAILY PERFORMANCE REPORT- PURE ARRAYS- {}'.format(now.strftime("%d-%m-%Y")))
    style = doc.styles['Normal']
    font = style.font

    font.name = 'Arial'
    font.size = Pt(16)
    title.bold = True
    title.underline = True

    doc.add_paragraph()
    doc.add_paragraph()
    
    for array in arrays:
        purearray = doc.add_paragraph().add_run(array)
        purearray.bold = True
        pic = r'D:\Temp\VM Pure Array Health check\Screenshots\{} [{}].png'.format(array, now.strftime("%d-%m-%Y"))
        doc.add_picture(pic, width=Inches(5))
        doc.add_paragraph()
        
        if arrays.index(array) %2 != 0:
            doc.add_page_break()
        
    report = 'D://Temp//VM Pure Array Health check//DAILY PERFORMANCE REPORT [{}].docx'.format(now.strftime("%d-%m-%Y"))
    
    doc.save(report)


if __name__ == "__main__":
    main()
    
    


